"""
API-friendly wrapper for contract generation
Supports dynamic contract data from the frontend
"""
from docxtpl import DocxTemplate
from datetime import datetime
import os
import sys
import subprocess
import time

# Configuration
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "templates", "pc-template.docx")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "output")
LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"


def generate_contract_with_data(contract_data):
    """
    Generate contract using provided data

    Args:
        contract_data: Dictionary containing contract information

    Returns:
        tuple: (docx_path, pdf_path) or (docx_path, None) if PDF fails
    """

    print(f"\n[API] Generating contract for: {contract_data.get('customerName')}")

    total_start = time.time()

    # Load template
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at {TEMPLATE_PATH}")

    doc = DocxTemplate(TEMPLATE_PATH)

    # Populate data
    docx_start = time.time()

    # Render template (excluding sites for now)
    data_without_sites = {k: v for k, v in contract_data.items() if k != 'sites'}
    doc.render(data_without_sites)

    # Save DOCX
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"contract_api_{timestamp}.docx"
    docx_path = os.path.join(OUTPUT_DIR, docx_filename)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(docx_path)

    # Add sites table if provided
    if contract_data.get('sites'):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        final_doc = Document(docx_path)

        # Find paragraph with "MVI = Move-In"
        target_idx = None
        for i, para in enumerate(final_doc.paragraphs):
            if "MVI = Move-In" in para.text:
                target_idx = i
                break

        if target_idx:
            sites = contract_data['sites']

            # Create table
            table = final_doc.add_table(rows=1 + len(sites), cols=4)
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'SERVICE ADDRESS'
            hdr_cells[1].text = 'ESI ID'
            hdr_cells[2].text = 'START OPTION\nMVI / SWI / REN'
            hdr_cells[3].text = 'START DATE\nif applicable'

            # Format header
            for cell in hdr_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                shading_xml = '<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))
                cell._element.get_or_add_tcPr().append(parse_xml(shading_xml))

            # Data rows
            for idx, site in enumerate(sites, 1):
                row_cells = table.rows[idx].cells

                address = f"{site.get('addressLine', '')}\n{site.get('city', '')} {site.get('state', '')} {site.get('zip', '')}"
                row_cells[0].text = address
                row_cells[1].text = site.get('esiid', '')

                start_option = site.get('startOption', '').lower()
                if start_option == 'move-in':
                    row_cells[2].text = "X Move-in\n  Switch"
                elif start_option == 'switch':
                    row_cells[2].text = "  Move-in\nX Switch"
                else:
                    row_cells[2].text = "  Move-in\n  Switch"

                row_cells[3].text = site.get('startDate', '')

            # Move table to correct position
            table_element = table._element
            final_doc.paragraphs[target_idx]._element.addnext(table_element)

        final_doc.save(docx_path)

    docx_time = time.time() - docx_start
    print(f"[API] DOCX generated in {docx_time:.3f}s: {docx_path}")

    # Convert to PDF
    pdf_filename = f"contract_api_{timestamp}.pdf"
    pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)

    pdf_start = time.time()
    success = convert_to_pdf_libreoffice(docx_path, pdf_path)
    pdf_time = time.time() - pdf_start

    if success:
        print(f"[API] PDF generated in {pdf_time:.3f}s: {pdf_path}")
    else:
        print(f"[API] PDF conversion failed")
        pdf_path = None

    total_time = time.time() - total_start
    print(f"[API] Total time: {total_time:.3f}s")

    return docx_path, pdf_path


def convert_to_pdf_libreoffice(docx_path, pdf_path):
    """Convert DOCX to PDF using LibreOffice headless mode"""

    if not os.path.exists(LIBREOFFICE_PATH):
        print(f"[WARN] LibreOffice not found at: {LIBREOFFICE_PATH}")
        return False

    try:
        output_dir = os.path.dirname(os.path.abspath(pdf_path))
        docx_abs_path = os.path.abspath(docx_path)

        cmd = [
            LIBREOFFICE_PATH,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_abs_path
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        return result.returncode == 0

    except Exception as e:
        print(f"[ERROR] PDF conversion failed: {e}")
        return False
