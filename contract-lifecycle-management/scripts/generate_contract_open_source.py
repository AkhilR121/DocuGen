"""
CLM POC - Open-Source Approach
Uses: docxtpl (Jinja2) + LibreOffice for PDF conversion
"""
from docxtpl import DocxTemplate
from datetime import datetime
import os
import sys
import subprocess
import time

# Configuration
TEMPLATE_PATH = "assets/templates/pc-template.docx"
OUTPUT_DIR = "output"
LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"

# Hardcoded contract data
CONTRACT_DATA = {
    "customerName": "Delray Oil Refining Company",
    "primaryContact": "primary contact name",
    "addressLine1": "901 louisiana st",
    "city": "Houston",
    "state": "Texas",
    "zip": "500018",
    "emailAddress": "akh@gmail.com",
    "billingAddressLine": "1901 louisiana st",
    "billingCity": "Houston",
    "billingState": "Texas",
    "billingZip": "77102",
    "language": "English",
    "contractPrice": "1.57",
    "startDate": "07/01/2026",
    "endDate": "06/30/2027",
    "contractDuration": "12 months",
    "sites": [
        {
            "esiid": "9999999",
            "startOption": "move-in",
            "startDate": "07/01/2026",
            "addressLine": "901 louisiana st",
            "city": "Houston",
            "state": "Texas",
            "zip": "77002",
        },
        {
            "esiid": "8888888",
            "startOption": "switch",
            "startDate": "07/01/2026",
            "addressLine": "901 louisiana st",
            "city": "Houston",
            "state": "Texas",
            "zip": "77002",
        },
        {
            "esiid": "7777777",
            "startOption": "switch",
            "startDate": "07/01/2026",
            "addressLine": "901 louisiana st",
            "city": "Houston",
            "state": "Texas",
            "zip": "77002",
        },
    ]
}


def generate_contract():
    """Generate contract using open-source approach"""

    print("=" * 80)
    print("CLM POC - OPEN-SOURCE APPROACH")
    print("Library: docxtpl (Jinja2) + LibreOffice")
    print("=" * 80)

    total_start = time.time()

    # Step 1: Load template
    print(f"\n[1/4] Loading template: {TEMPLATE_PATH}")
    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERROR: Template not found at {TEMPLATE_PATH}")
        sys.exit(1)

    try:
        doc = DocxTemplate(TEMPLATE_PATH)
        print("[OK] Template loaded")
    except Exception as e:
        print(f"ERROR: {e}")
        sys.exit(1)

    # Step 2: Populate data and save DOCX
    print(f"\n[2/4] Populating data and generating DOCX...")
    print(f"      Customer: {CONTRACT_DATA['customerName']}")
    print(f"      Location: {CONTRACT_DATA['city']}, {CONTRACT_DATA['state']}")
    print(f"      Period: {CONTRACT_DATA['startDate']} - {CONTRACT_DATA['endDate']}")

    docx_start = time.time()

    try:
        # Render template with data (excluding sites for now)
        data_without_sites = {k: v for k, v in CONTRACT_DATA.items() if k != 'sites'}
        doc.render(data_without_sites)

        # Save DOCX temporarily
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"contract_opensource_{timestamp}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)

        os.makedirs(OUTPUT_DIR, exist_ok=True)
        doc.save(docx_path)

        # Now add sites table using python-docx
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        final_doc = Document(docx_path)

        # Find paragraph with "MVI = Move-In"
        target_idx = None
        for i, para in enumerate(final_doc.paragraphs):
            if "MVI = Move-In" in para.text:
                target_idx = i
                break

        if target_idx and CONTRACT_DATA.get('sites'):
            # Insert table after MVI paragraph
            sites = CONTRACT_DATA['sites']

            # Create table (4 columns, 1 header + N data rows)
            table = final_doc.add_table(rows=1 + len(sites), cols=4)
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'SERVICE ADDRESS'
            hdr_cells[1].text = 'ESI ID'
            hdr_cells[2].text = 'START OPTION\nMVI / SWI / REN'
            hdr_cells[3].text = 'START DATE\nif applicable'

            # Format header
            for cell in hdr_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                # Add shading
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                shading_xml = '<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))
                cell._element.get_or_add_tcPr().append(parse_xml(shading_xml))

            # Data rows
            for idx, site in enumerate(sites, 1):
                row_cells = table.rows[idx].cells

                # Service Address
                address = f"{site.get('addressLine', '')}\n{site.get('city', '')} {site.get('state', '')} {site.get('zip', '')}"
                row_cells[0].text = address

                # ESI ID
                row_cells[1].text = site.get('esiid', '')

                # Start Option
                start_option = site.get('startOption', '').lower()
                if start_option == 'move-in':
                    row_cells[2].text = "X Move-in\n  Switch"
                elif start_option == 'switch':
                    row_cells[2].text = "  Move-in\nX Switch"
                else:
                    row_cells[2].text = "  Move-in\n  Switch"

                # Start Date
                row_cells[3].text = site.get('startDate', '')

            # Move table to correct position (after MVI paragraph)
            table_element = table._element
            final_doc.paragraphs[target_idx]._element.addnext(table_element)

            print(f"      Inserted sites table with {len(sites)} rows")

        # Save final document
        final_doc.save(docx_path)

        docx_time = time.time() - docx_start

        print(f"[OK] DOCX generated: {docx_path}")
        print(f"      Time taken: {docx_time:.3f} seconds")

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Step 3: Convert to PDF using LibreOffice
    pdf_filename = f"contract_opensource_{timestamp}.pdf"
    pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)

    print(f"\n[3/4] Converting DOCX to PDF using LibreOffice...")

    pdf_start = time.time()
    success = convert_to_pdf_libreoffice(docx_path, pdf_path)
    pdf_time = time.time() - pdf_start

    if success:
        print(f"[OK] PDF generated: {pdf_path}")
        print(f"      Time taken: {pdf_time:.3f} seconds")
    else:
        print("[FAILED] PDF conversion unsuccessful")
        print("\nLibreOffice Installation:")
        print("  1. Download: https://www.libreoffice.org/download/")
        print("  2. Install to default location")
        print("  3. Or update LIBREOFFICE_PATH in script")
        pdf_path = None

    # Summary
    total_time = time.time() - total_start

    print("\n" + "=" * 80)
    print("GENERATION COMPLETE - OPEN-SOURCE APPROACH")
    print("=" * 80)
    print(f"\nPerformance Metrics:")
    print(f"  DOCX Generation:    {docx_time:.3f} seconds")
    if success:
        print(f"  PDF Conversion:     {pdf_time:.3f} seconds")
        print(f"  Total Time:         {total_time:.3f} seconds")
    else:
        print(f"  PDF Conversion:     FAILED")
        print(f"  Total Time (DOCX):  {docx_time:.3f} seconds")

    print(f"\nGenerated Files:")
    print(f"  DOCX: {os.path.abspath(docx_path)}")
    if success:
        print(f"  PDF:  {os.path.abspath(pdf_path)}")
    print()

    return docx_path, pdf_path


def convert_to_pdf_libreoffice(docx_path, pdf_path):
    """Convert DOCX to PDF using LibreOffice headless mode"""

    if not os.path.exists(LIBREOFFICE_PATH):
        print(f"[ERROR] LibreOffice not found at: {LIBREOFFICE_PATH}")
        return False

    try:
        output_dir = os.path.dirname(os.path.abspath(pdf_path))
        docx_abs_path = os.path.abspath(docx_path)

        cmd = [
            LIBREOFFICE_PATH,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_abs_path
        ]

        print(f"      Executing LibreOffice conversion...")
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode == 0:
            return True
        else:
            print(f"[ERROR] LibreOffice failed: {result.stderr}")
            return False

    except subprocess.TimeoutExpired:
        print("[ERROR] Conversion timeout (>60s)")
        return False
    except Exception as e:
        print(f"[ERROR] {e}")
        return False


if __name__ == "__main__":
    generate_contract()
